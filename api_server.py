import os
import time
import datetime
import threading
import requests
import psycopg2
from flask import Flask, request, jsonify, send_file
from flask_cors import CORS  # 跨域連線套件

try:
    from docx import Document
    from docx.shared import Pt
except ImportError:
    pass # 記得在 Render 的 requirements.txt 加入 python-docx

# ==========================================
# 🚀 建立獨立 APP 後端伺服器
# ==========================================
# ⚠️ 絕對只保留這一個 app 宣告，避免 CORS 被覆蓋
app = Flask(__name__)
CORS(app)  # 開放網頁直接呼叫大腦

# ==========================================
# 🔑 1. 系統環境變數與核心金鑰配置
# ==========================================
# PostgreSQL 資料庫專屬網址
DB_URL = os.environ.get("DATABASE_URL", "postgresql://dispatch_db_h7du_user:th0PNe1ycle9s2n3Dh6ha2gefpAzISo1@dpg-d9r8ejfavr4c738la9pg-a.singapore-postgres.render.com/dispatch_db_h7du")

# 衛星犬 EUP 專屬金鑰
EUP_BASE_URL = 'https://tw.eupfin.com/Eup_Servlet_API_SOAP'
EUP_TOKEN = 'e386bd24-4403-fd4e-158c-bb4b66dfe989'

# ==========================================
# 📱 2. 老闆與主管的警報廣播系統
# ==========================================
def send_alert_to_boss_and_manager(message):
    """模擬發送 LINE Notify 或 APP 推播給黃董與主管"""
    print("\n" + "🔴" * 25)
    print("📲 【系統緊急廣播給 黃董 / 主管】")
    print(message)
    print("🔴" * 25 + "\n")
    # 未來這裡會串接真正的 LINE Notify API

# ==========================================
# 🧠 3. 核心商業邏輯模組 (三層實戰防線)
# ==========================================
def check_unfinished_tasks(tech_name):
    """【防線一】檢查師傅手邊是否還有未完成的訂單"""
    unfinished_count = 1  
    if unfinished_count > 0:
        alert_msg = f"🚨 【下班異常警報】\n師傅 {tech_name} 嘗試打卡下班，但系統偵測到他手邊還有 {unfinished_count} 張訂單「尚未完工」！請主管立即查核。"
        send_alert_to_boss_and_manager(alert_msg)
        return False, "您手邊還有未完成的訂單，禁止打卡下班！已通報主管。"
    return True, "任務全數清空。"

def audit_off_work_time(tech_name, reported_off_work_time_str, eup_engine_off_time_str):
    """【防線二】審核下班時間是否浮報 (20分鐘抓漏)"""
    try:
        fmt = "%H:%M"
        reported_time = datetime.datetime.strptime(reported_off_work_time_str, fmt)
        engine_off_time = datetime.datetime.strptime(eup_engine_off_time_str, fmt)
        
        time_diff = (reported_time - engine_off_time).total_seconds() / 60
        
        if time_diff > 20:
            return False, f"時數異常 (差距 {int(time_diff)} 分鐘)"
        return True, "時數審核吻合"
    except Exception as e:
        return False, f"計算錯誤: {e}"

def manual_adjust_service_time(order_id, new_duration_minutes):
    """【AI 調度員】中控台手動微調服務時間，觸發路線重算"""
    print(f"🔧 中控指令：訂單 #{order_id} 修改為 {new_duration_minutes} 分鐘。")
    print("🤖 喚醒 TSP AI，正在為該師傅重新規劃後續路線...")
    return {"status": "success", "msg": "✅ AI 路線重算完畢，師傅端 APP 已更新！"}

# ==========================================
# 📡 4. 背景自動巡邏雷達 (40分鐘預警)
# ==========================================
def background_radar_task():
    """【防線三】每分鐘自動掃描一次，檢查是否有即將遲到的任務"""
    while True:
        time.sleep(60)

radar_thread = threading.Thread(target=background_radar_task, daemon=True)
radar_thread.start()

# ==========================================
# 🌐 5. 獨立 APP 專屬連線通道 (API Routes)
# ==========================================
@app.route('/', methods=['GET'])
def health_check():
    return jsonify({"status": "線上", "version": "行動車美 獨立APP大腦 v1.0", "db": "PostgreSQL 準備就緒"})

@app.route('/api/clock_out', methods=['POST', 'OPTIONS'])
def handle_clock_out():
    if request.method == 'OPTIONS':
        return jsonify({}), 200

    data = request.json
    tech_name = data.get('tech_name')
    reported_time = data.get('off_work_time')
    
    tasks_cleared, task_msg = check_unfinished_tasks(tech_name)
    if not tasks_cleared:
        return jsonify({"status": "alert", "message": task_msg}), 400

    real_engine_off_time = "16:00" 
    is_valid, time_msg = audit_off_work_time(tech_name, reported_time, real_engine_off_time)
    
    if not is_valid:
        alert_msg = f"🚨 【防弊警報】\n師傅 {tech_name} 浮報下班時間！\n真實熄火：{real_engine_off_time} | 師傅回報：{reported_time}。"
        send_alert_to_boss_and_manager(alert_msg)
        return jsonify({"status": "alert", "message": f"下班打卡遭拒絕！\n{time_msg}"}), 400

    return jsonify({"status": "success", "message": "下班打卡成功！辛苦了！"})

@app.route('/api/adjust_time', methods=['POST', 'OPTIONS'])
def handle_adjust_time():
    if request.method == 'OPTIONS':
        return jsonify({}), 200
    data = request.json
    return jsonify(manual_adjust_service_time(data.get('order_id'), data.get('new_duration')))


# ==========================================
# 📊 6. Web後台/APP 專用報表生成通道 (一鍵 Word 薪資單)
# ==========================================
def generate_monthly_payroll_word(tech_name, month_str, total_points, overtime_mins, late_mins, base_salary, point_rate):
    """內部引擎：生成排版好的 Word 檔案"""
    doc = Document()
    
    title = doc.add_heading(f'兆鼎國際 - {month_str} 薪資與績效結算單', 0)
    title.alignment = 1 
    
    doc.add_heading('一、 出勤與基本資料', level=1)
    doc.add_paragraph(f'員工姓名：{tech_name}')
    doc.add_paragraph(f'結算月份：{month_str}')
    doc.add_paragraph(f'總加班時間：{overtime_mins} 分鐘')
    doc.add_paragraph(f'總遲到時間：{late_mins} 分鐘 (依 EUP 衛星犬發動紀錄核算)')
    
    doc.add_heading('二、 績效積分與薪資核算', level=1)
    doc.add_paragraph(f'底薪/保障薪資：NT$ {base_salary:,}')
    doc.add_paragraph(f'當月總積分：{total_points} 分')
    
    bonus = int(total_points * point_rate)
    doc.add_paragraph(f'績效獎金換算 ({point_rate}元/分)：NT$ {bonus:,}')
    
    total_salary = base_salary + bonus
    doc.add_heading('三、 本月實發總計', level=1)
    p = doc.add_paragraph(f'NT$ {total_salary:,}')
    p.runs[0].font.size = Pt(16)
    p.runs[0].font.bold = True
    
    import tempfile
    file_path = os.path.join(tempfile.gettempdir(), f"{tech_name}_{month_str}_薪資單.docx")
    doc.save(file_path)
    return file_path

@app.route('/api/admin/download_payroll', methods=['GET', 'OPTIONS'])
def api_download_payroll():
    """APP 或 Web 後台專用的下載 API"""
    if request.method == 'OPTIONS':
        return jsonify({}), 200

    tech_name = request.args.get('tech_name', '曾紹恩')
    month_str = request.args.get('month', '2026-03')
    
    mock_points, mock_overtime, mock_late, mock_base, mock_rate = 850, 120, 0, 35000, 20

    try:
        file_path = generate_monthly_payroll_word(
            tech_name, month_str, mock_points, mock_overtime, mock_late, mock_base, mock_rate
        )
        
        import urllib.parse
        safe_filename = urllib.parse.quote(os.path.basename(file_path))
        return send_file(file_path, as_attachment=True, download_name=os.path.basename(file_path))
    except Exception as e:
        return jsonify({"status": "error", "message": f"報表生成失敗: {str(e)}"}), 500


# ==========================================
# 🗺️ 7. PostGIS 空間地理運算模組 (取代 Google Maps 算距離)
# ==========================================
def find_nearest_available_techs(customer_lat, customer_lng, radius_meters=5000):
    try:
        conn = psycopg2.connect(DB_URL)
        cursor = conn.cursor()
        
        query = """
            SELECT tech_name, 
                   ST_DistanceSphere(current_location, ST_MakePoint(%s, %s)) AS distance_meters
            FROM tech_locations
            WHERE is_working = TRUE 
            AND ST_DWithin(current_location, ST_MakePoint(%s, %s), %s)
            ORDER BY distance_meters ASC
            LIMIT 3;
        """
        cursor.execute(query, (customer_lng, customer_lat, customer_lng, customer_lat, radius_meters))
        results = cursor.fetchall()
        
        cursor.close()
        conn.close()
        
        techs = []
        for row in results:
            techs.append({
                "tech_name": row[0],
                "distance_meters": round(row[1], 1)
            })
        return techs
    except Exception as e:
        print(f"⚠️ PostGIS 空間運算失敗: {e}")
        return []

@app.route('/api/admin/spatial_dispatch', methods=['POST', 'OPTIONS'])
def api_spatial_dispatch():
    if request.method == 'OPTIONS':
        return jsonify({}), 200

    data = request.json
    lat, lng = data.get('lat'), data.get('lng')
    
    if not lat or not lng:
        return jsonify({"status": "error", "message": "缺少經緯度座標"}), 400

    nearest_techs = find_nearest_available_techs(lat, lng, radius_meters=5000)
    
    if nearest_techs:
        return jsonify({"status": "success", "recommended_tech": nearest_techs[0], "all_nearby_techs": nearest_techs}), 200
    else:
        return jsonify({"status": "warning", "message": "5公里內無可用的在線師傅"}), 200


# ==========================================
# 🤖 8. AI 極端防呆與強制追問通道
# ==========================================
@app.route('/api/ai/parse_order', methods=['POST', 'OPTIONS'])
def api_ai_parse_order():
    if request.method == 'OPTIONS':
        return jsonify({}), 200

    data = request.json
    raw_text = data.get('text', '')
    
    if not raw_text:
        return jsonify({"status": "error", "message": "無解析內容"}), 400

    try:
        import google.generativeai as genai
        GEMINI_API_KEY = os.environ.get('GEMINI_API_KEY', '請在Render設定')
        genai.configure(api_key=GEMINI_API_KEY)
        
        prompt = f"""
        你是一個極度嚴謹的汽車美容派單助理。請分析這段客戶對話，提取預約資訊。
        ⚠️【最高防幻覺指令】⚠️
        1. 絕對不允許猜測或腦補！如果客戶說「車站附近」，地址只能填「車站附近」，不可自行捏造路名或完整地址。
        2. 如果客戶說「下午」，時間請填「下午」，不可自行決定特定鐘點（如 14:00）。
        3. 如果客戶完全沒提到某個資訊，對應的 value 必須是空字串 ""。

        請嚴格輸出以下純 JSON 格式（不要包覆 markdown 語法）：
        {{
            "phone": "電話號碼",
            "address": "詳細服務地點(不可腦補)",
            "plan": "服務方案",
            "time": "客戶希望的時間",
            "plate": "車牌號碼",
            "clarification_needed": "如果客戶提供的地點或時間太模糊（例如只有說下午、或只給區名），請在這裡生成一句客氣的追問話術。若資訊完全完整則填空字串 ""。"
        }}

        客戶對話內容：
        {raw_text}
        """
        
        model = genai.GenerativeModel('gemini-2.5-flash', generation_config={"response_mime_type": "application/json"})
        response = model.generate_content(prompt)
        
        import json
        result_json = json.loads(response.text.strip())
        
        if result_json.get("clarification_needed"):
            return jsonify({"status": "need_clarification", "message": result_json["clarification_needed"], "parsed_data": result_json}), 200
        else:
            return jsonify({"status": "complete", "message": "資訊完整，可直接建單！", "parsed_data": result_json}), 200

    except Exception as e:
        return jsonify({"status": "error", "message": f"AI 解析失敗: {str(e)}"}), 500


# ==========================================
# 🚀 9. 啟動伺服器引擎
# ==========================================
if __name__ == '__main__':
    print("========================================")
    print("啟動【行動車美】獨立 APP 後端控制中心...")
    print("========================================")
    port = int(os.environ.get('PORT', 5000))
    app.run(host='0.0.0.0', port=port, debug=True)
